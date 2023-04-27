import docx

document = docx.Document()

document.add_paragraph('Hello World', 'Title')
document.add_paragraph('by Brian', 'Heading 1')
document.add_paragraph('This is a Word document created by Python using the python-docx library.')
document.add_paragraph('Python programming is challenging!', 'Intense Quote')
document.add_paragraph('This paragraph has some ')
last_paragraph = document.paragraphs[-1]
last_paragraph.add_run('bold text')
last_run = last_paragraph.runs[-1]
last_run.bold = True
last_run.underline = True
last_paragraph.add_run(' in the middle.')
document.save('hello_word_style.docx')